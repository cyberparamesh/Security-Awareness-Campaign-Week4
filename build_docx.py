import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_field(p, text):
    run = p.add_run()
    r = run._r
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> %s </w:instrText>' % (nsdecls('w'), text))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_formatted_runs(p, text, base_font_size=10.5, base_color=RGBColor(0x2D, 0x37, 0x48), is_italic=False):
    tokens = re.split(r'(\[.*?\]\(.*?\)|\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for token in tokens:
        if not token:
            continue
        link_match = re.match(r'^\[(.*?)\]\((.*?)\)$', token)
        if link_match:
            link_text, link_url = link_match.groups()
            run = p.add_run(f"{link_text} ({link_url})")
            run.font.name = 'Segoe UI'
            run.font.size = Pt(base_font_size)
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            run.font.underline = True
            continue

        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            run = p.add_run(token[2:-2])
            run.font.name = 'Segoe UI'
            run.font.size = Pt(base_font_size)
            run.font.bold = True
            run.font.color.rgb = base_color
            if is_italic:
                run.font.italic = True
            continue

        if token.startswith('`') and token.endswith('`') and len(token) >= 2:
            run = p.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(base_font_size - 1)
            run.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)
            continue

        if token.startswith('*') and token.endswith('*') and len(token) >= 2 and not token.startswith('**'):
            run = p.add_run(token[1:-1])
            run.font.name = 'Segoe UI'
            run.font.size = Pt(base_font_size)
            run.font.italic = True
            run.font.color.rgb = base_color
            continue

        run = p.add_run(token)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(base_font_size)
        run.font.color.rgb = base_color
        if is_italic:
            run.font.italic = True

def build_docx():
    md_path = r"d:\cyber security analysis\Security-Awareness-Campaign-Week4\report\Week4_Security_Awareness_Campaign.md"
    docx_path = r"d:\cyber security analysis\Security-Awareness-Campaign-Week4\Week4_Security_Awareness_Campaign.docx"
    docx_report_path = r"d:\cyber security analysis\Security-Awareness-Campaign-Week4\report\Week4_Security_Awareness_Campaign.docx"

    doc = Document()

    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("YuvaIntern Cyber Security Analyst Internship  |  Week 4 Report")
    hrun.font.name = "Segoe UI"
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    frun1 = fp.add_run("Security Awareness Campaign Plan  |  Page ")
    frun1.font.name = "Segoe UI"
    frun1.font.size = Pt(9)
    frun1.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    add_field(fp, "PAGE")
    frun2 = fp.add_run(" of ")
    frun2.font.name = "Segoe UI"
    frun2.font.size = Pt(9)
    frun2.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    add_field(fp, "NUMPAGES")

    # --- TITLE PAGE ---
    t_space = doc.add_paragraph()
    t_space.paragraph_format.space_before = Pt(36)

    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_after = Pt(12)
    r_badge = p_badge.add_run("YUVAINTERN CYBER SECURITY ANALYST INTERNSHIP PROGRAM")
    r_badge.font.name = "Segoe UI"
    r_badge.font.size = Pt(10)
    r_badge.font.bold = True
    r_badge.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(8)
    p_title.paragraph_format.line_spacing = 1.15
    r_title = p_title.add_run("Enterprise Security Awareness Campaign Plan")
    r_title.font.name = "Segoe UI"
    r_title.font.size = Pt(24)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(36)
    r_sub = p_sub.add_run("Week 4 – 12-Week Human Risk Management Framework & Culture Change Strategy")
    r_sub.font.name = "Segoe UI"
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    div_table = doc.add_table(rows=1, cols=1)
    div_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    div_table.autofit = False
    div_table.columns[0].width = Inches(6.5)
    cell = div_table.cell(0, 0)
    set_cell_background(cell, "1B365D")
    set_cell_margins(cell, top=20, bottom=20, left=0, right=0)
    p_div = cell.paragraphs[0]
    p_div.paragraph_format.space_before = Pt(0)
    p_div.paragraph_format.space_after = Pt(0)

    p_mspace = doc.add_paragraph()
    p_mspace.paragraph_format.space_before = Pt(40)

    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    meta_table.columns[0].width = Inches(1.8)
    meta_table.columns[1].width = Inches(4.7)

    metadata_items = [
        ("Intern Name:", "Parameshwaran"),
        ("Role / Track:", "Cyber Security Analyst – Internship Week 4"),
        ("Target Profile:", "Apex Global Technologies (Hypothetical Enterprise Profile)"),
        ("Framework Alignment:", "NIST SP 800-50, NIST SP 800-53 AT Controls, CISA Secure Our World"),
        ("Submission Date:", "August 2026")
    ]

    for idx, (label, val) in enumerate(metadata_items):
        cell_lbl = meta_table.cell(idx, 0)
        cell_val = meta_table.cell(idx, 1)

        set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
        set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)

        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(2)
        r_l = p_lbl.add_run(label)
        r_l.font.name = "Segoe UI"
        r_l.font.size = Pt(10.5)
        r_l.font.bold = True
        r_l.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

        p_v = cell_val.paragraphs[0]
        p_v.paragraph_format.space_after = Pt(2)
        r_v = p_v.add_run(val)
        r_v.font.name = "Segoe UI"
        r_v.font.size = Pt(10.5)
        r_v.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    doc.add_page_break()

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_buffer = []
    in_table = False
    table_lines = []

    def flush_table(lines_list):
        if not lines_list:
            return
        header_row = [c.strip() for c in lines_list[0].strip().strip('|').split('|')]
        data_rows = []
        for line in lines_list[2:]:
            if '|' in line:
                row = [c.strip() for c in line.strip().strip('|').split('|')]
                data_rows.append(row)

        tbl = doc.add_table(rows=1 + len(data_rows), cols=len(header_row))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False

        set_table_borders(tbl)

        for idx, col_text in enumerate(header_row):
            c = tbl.cell(0, idx)
            set_cell_background(c, "1B365D")
            set_cell_margins(c, top=100, bottom=100, left=120, right=120)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_runs(p, col_text, base_font_size=10, base_color=RGBColor(0xFF, 0xFF, 0xFF))
            if p.runs:
                p.runs[0].font.bold = True

        for r_idx, r_data in enumerate(data_rows):
            fill_bg = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell_text in enumerate(r_data):
                if c_idx < len(header_row):
                    c = tbl.cell(r_idx + 1, c_idx)
                    set_cell_background(c, fill_bg)
                    set_cell_margins(c, top=90, bottom=90, left=120, right=120)
                    p = c.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    add_formatted_runs(p, cell_text, base_font_size=9.5, base_color=RGBColor(0x2D, 0x37, 0x48))

        p_sp = doc.add_paragraph()
        p_sp.paragraph_format.space_before = Pt(0)
        p_sp.paragraph_format.space_after = Pt(8)

    def flush_code_block(code_lines):
        if not code_lines:
            return
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        tbl.columns[0].width = Inches(6.5)
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F7FAFC")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        
        tcPr = cell._element.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/><w:left w:val="single" w:sz="12" w:space="0" w:color="1B365D"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/><w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/></w:tcBorders>')
        tcPr.append(borders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05

        full_code = "".join(code_lines)
        run = p.add_run(full_code.rstrip())
        run.font.name = "Consolas"
        run.font.size = Pt(9.0)
        run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

        p_sp = doc.add_paragraph()
        p_sp.paragraph_format.space_after = Pt(6)

    skip_header_lines = True
    header_line_count = 0

    for line in lines:
        if skip_header_lines:
            if line.strip() == "---":
                header_line_count += 1
                if header_line_count >= 1:
                    skip_header_lines = False
            continue

        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                flush_code_block(code_buffer)
                code_buffer = []
            else:
                if in_table:
                    in_table = False
                    flush_table(table_lines)
                    table_lines = []
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if '|' in line and (line.strip().startswith('|') or line.strip().endswith('|')):
            if not in_table:
                in_table = True
            table_lines.append(line)
            continue
        else:
            if in_table:
                in_table = False
                flush_table(table_lines)
                table_lines = []

        stripped = line.strip()
        if not stripped:
            continue

        if stripped == "---":
            p_div = doc.add_paragraph()
            p_div.paragraph_format.space_before = Pt(8)
            p_div.paragraph_format.space_after = Pt(12)
            r = p_div.add_run("_________________________________________________________________________________")
            r.font.name = "Segoe UI"
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
            continue

        if stripped.startswith("## "):
            h_text = stripped[3:].strip()
            p_h1 = doc.add_paragraph()
            p_h1.paragraph_format.space_before = Pt(18)
            p_h1.paragraph_format.space_after = Pt(6)
            p_h1.paragraph_format.keep_with_next = True
            add_formatted_runs(p_h1, h_text, base_font_size=15, base_color=RGBColor(0x1B, 0x36, 0x5D))
            if p_h1.runs:
                p_h1.runs[0].font.bold = True
            continue

        if stripped.startswith("### "):
            h_text = stripped[4:].strip()
            p_h2 = doc.add_paragraph()
            p_h2.paragraph_format.space_before = Pt(14)
            p_h2.paragraph_format.space_after = Pt(4)
            p_h2.paragraph_format.keep_with_next = True
            add_formatted_runs(p_h2, h_text, base_font_size=12.5, base_color=RGBColor(0x2B, 0x6C, 0xB0))
            if p_h2.runs:
                p_h2.runs[0].font.bold = True
            continue

        if stripped.startswith("* ") or stripped.startswith("- "):
            p_list = doc.add_paragraph()
            p_list.paragraph_format.left_indent = Inches(0.25)
            p_list.paragraph_format.space_before = Pt(2)
            p_list.paragraph_format.space_after = Pt(3)
            r_bullet = p_list.add_run("•  ")
            r_bullet.font.name = "Segoe UI"
            r_bullet.font.size = Pt(10.5)
            r_bullet.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            add_formatted_runs(p_list, stripped[2:], base_font_size=10.5)
            continue

        num_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if num_match:
            num, item_text = num_match.groups()
            p_list = doc.add_paragraph()
            p_list.paragraph_format.left_indent = Inches(0.25)
            p_list.paragraph_format.space_before = Pt(3)
            p_list.paragraph_format.space_after = Pt(4)
            r_num = p_list.add_run(f"{num}.  ")
            r_num.font.name = "Segoe UI"
            r_num.font.bold = True
            r_num.font.size = Pt(10.5)
            r_num.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            add_formatted_runs(p_list, item_text, base_font_size=10.5)
            continue

        p_para = doc.add_paragraph()
        p_para.paragraph_format.space_before = Pt(3)
        p_para.paragraph_format.space_after = Pt(6)
        p_para.paragraph_format.line_spacing = 1.15
        
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            add_formatted_runs(p_para, stripped[1:-1], base_font_size=10.5, is_italic=True)
        else:
            add_formatted_runs(p_para, stripped, base_font_size=10.5)

    doc.save(docx_path)
    doc.save(docx_report_path)
    print(f"Successfully generated DOCX files:\n - {docx_path}\n - {docx_report_path}")

if __name__ == "__main__":
    build_docx()
